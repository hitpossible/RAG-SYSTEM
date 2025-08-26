import os
from typing import List, Dict
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd

class DocumentLoader:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def load_pdf(self, file_path: str) -> List[Document]:
        """Load and process PDF files"""
        documents = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                
                chunks = self.text_splitter.split_text(text)
                for i, chunk in enumerate(chunks):
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "page": i,
                            "type": "pdf"
                        }
                    )
                    documents.append(doc)
        except Exception as e:
            print(f"Error loading PDF {file_path}: {e}")
        
        return documents
    
    def load_docx(self, file_path: str) -> List[Document]:
        """Load and process DOCX files"""
        documents = []
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            chunks = self.text_splitter.split_text(text)
            for i, chunk in enumerate(chunks):
                document = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk": i,
                        "type": "docx"
                    }
                )
                documents.append(document)
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {e}")
        
        return documents
    
    def load_txt(self, file_path: str) -> List[Document]:
        """Load and process TXT files"""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            chunks = self.text_splitter.split_text(text)
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk": i,
                        "type": "txt"
                    }
                )
                documents.append(doc)
        except Exception as e:
            print(f"Error loading TXT {file_path}: {e}")
        
        return documents
    
    def load_directory(self, directory_path: str) -> List[Document]:
        """Load all supported documents from directory"""
        all_documents = []
        
        if not os.path.exists(directory_path):
            print(f"Directory {directory_path} does not exist")
            return all_documents
        
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            if filename.lower().endswith('.pdf'):
                all_documents.extend(self.load_pdf(file_path))
            elif filename.lower().endswith('.docx'):
                all_documents.extend(self.load_docx(file_path))
            elif filename.lower().endswith('.txt'):
                all_documents.extend(self.load_txt(file_path))
        
        print(f"Loaded {len(all_documents)} document chunks")
        return all_documents