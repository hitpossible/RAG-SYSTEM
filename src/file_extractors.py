# src/file_extractors.py
from io import BytesIO
from typing import Tuple

def extract_text_from_pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as e:
        raise RuntimeError("Missing dependency: pdfplumber. pip install pdfplumber") from e

    text_parts = []
    with pdfplumber.open(BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
    return "\n\n".join(text_parts).strip()

def extract_text_from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("Missing dependency: python-docx. pip install python-docx") from e

    doc = Document(BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # ตาราง (ถ้ามี)
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()

def extract_text_from_excel(data: bytes) -> str:
    try:
        import pandas as pd
    except ImportError as e:
        raise RuntimeError("Missing dependency: pandas (and openpyxl/xlrd). "
                           "pip install pandas openpyxl xlrd") from e

    # อ่านทุกชีต -> ต่อเป็นข้อความเดียว
    buf = BytesIO(data)
    # engine เลือกตามนามสกุลโดย pandas
    sheets = pd.read_excel(buf, sheet_name=None, dtype=str)
    out_parts = []
    for name, df in sheets.items():
        # แปลง NaN -> "" แล้วพิมพ์เป็นข้อความแบบตาราง
        df = df.fillna("")
        out_parts.append(f"# Sheet: {name}")
        # ใช้ to_csv เพื่อคุมฟอร์แมตอ่านง่าย
        out_parts.append(df.to_csv(index=False))
    return "\n\n".join(out_parts).strip()

def sniff_and_extract(filename: str, content_type: str, data: bytes) -> Tuple[str, str]:
    """
    คืนค่า: (extracted_text, kind)  โดย kind ∈ {"pdf","docx","excel"}
    """
    name = (filename or "").lower()
    ct = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ct:
        return extract_text_from_pdf(data), "pdf"
    if name.endswith(".docx") or "officedocument.wordprocessingml.document" in ct:
        return extract_text_from_docx(data), "docx"
    if name.endswith(".xlsx") or name.endswith(".xls") or "spreadsheetml" in ct or "excel" in ct:
        return extract_text_from_excel(data), "excel"

    raise ValueError(f"Unsupported file type: {filename} ({content_type})")
